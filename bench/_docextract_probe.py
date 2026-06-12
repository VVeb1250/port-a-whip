"""Capability probe for the doc-extract set (markitdown): does structure survive?

Generates a docx + xlsx fixture WITH structure (headings, a table, a sheet of rows),
converts each via markitdown, and checks the markdown preserves headings + tabular
shape. Capability-class (no token delta claim) — the gate is "binary file becomes
greppable structured md", per the set's neutral profile.

Usage: py bench/_docextract_probe.py
"""

from __future__ import annotations

import os
import shutil
import subprocess
from pathlib import Path

OUT = Path("bench/out")
OUT.mkdir(parents=True, exist_ok=True)
DOCX = OUT / "de_report.docx"
XLSX = OUT / "de_sheet.xlsx"


def make_docx() -> None:
    from docx import Document

    d = Document()
    d.add_heading("Quarterly Report", level=1)
    d.add_heading("Summary", level=2)
    d.add_paragraph("Revenue rose across all regions this quarter.")
    d.add_heading("Regional Breakdown", level=2)
    t = d.add_table(rows=1, cols=3)
    t.rows[0].cells[0].text = "Region"
    t.rows[0].cells[1].text = "Revenue"
    t.rows[0].cells[2].text = "Growth"
    for region, rev, gr in [("North", "120k", "8%"), ("South", "95k", "5%"), ("West", "210k", "12%")]:
        c = t.add_row().cells
        c[0].text, c[1].text, c[2].text = region, rev, gr
    d.save(DOCX)


def make_xlsx() -> None:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "sales"
    ws.append(["product", "units", "price"])
    for i in range(1, 11):
        ws.append([f"SKU-{i:02d}", i * 7, round(i * 3.5, 2)])
    wb.save(XLSX)


def md(path: Path) -> str:
    exe = shutil.which("markitdown") or "markitdown"
    p = subprocess.run([exe, str(path)], capture_output=True, text=True, encoding="utf-8")
    return (p.stdout or "") + (("\n[stderr] " + p.stderr) if p.stderr else "")


def main() -> int:
    make_docx()
    make_xlsx()
    print("# doc-extract capability probe (markitdown)\n")

    docx_md = md(DOCX)
    print("## docx -> md")
    print(docx_md.strip()[:600])
    docx_ok = "# Quarterly Report" in docx_md or "Quarterly Report" in docx_md
    table_ok = "Region" in docx_md and "Revenue" in docx_md and "|" in docx_md
    head_ok = "Summary" in docx_md and "Regional Breakdown" in docx_md
    print(f"\n  heading title present: {docx_ok} | sub-headings: {head_ok} | table-as-md: {table_ok}\n")

    xlsx_md = md(XLSX)
    print("## xlsx -> md")
    print(xlsx_md.strip()[:600])
    rows_ok = "SKU-01" in xlsx_md and "SKU-10" in xlsx_md
    cols_ok = "product" in xlsx_md and "units" in xlsx_md and "price" in xlsx_md
    print(f"\n  header cols present: {cols_ok} | all rows present: {rows_ok}\n")

    verdict = all([docx_ok, head_ok, table_ok, rows_ok, cols_ok])
    print(f"VERDICT: structure-preserving = {verdict}")
    return 0 if verdict else 1


if __name__ == "__main__":
    raise SystemExit(main())
